"""
REPORT EXPORT ENGINE — Multi-Format Report Generation
=======================================================
Generate reports in Markdown, PDF, and DOCX formats.

Usage:
    from tools.report_export import ReportExporter

    exporter = ReportExporter()
    md = exporter.to_markdown(report_data)
    pdf = exporter.to_pdf(report_data)
    docx = exporter.to_docx(report_data)
"""

import json
import os
from dataclasses import dataclass
from datetime import datetime
from typing import Dict, List, Any, Optional


REPORT_FORMATS = {
    "md": "md",
    "markdown": "md",
    "pdf": "pdf",
    "docx": "docx",
}


class ReportExportError(RuntimeError):
    """Typed error raised when a report cannot be exported safely."""

    code = "report_export_error"
    status_code = 503

    def __init__(self, message: str, *, details: Optional[Dict[str, Any]] = None):
        super().__init__(message)
        self.details = details or {}

    def as_detail(self) -> Dict[str, Any]:
        return {
            "code": self.code,
            "message": str(self),
            **self.details,
        }


class ReportNotReadyError(ReportExportError):
    code = "report_not_ready"
    status_code = 409


class ReportUnavailableError(ReportExportError):
    code = "report_unavailable"
    status_code = 503


class UnsupportedReportFormatError(ReportExportError):
    code = "unsupported_report_format"
    status_code = 400


class ReportDependencyError(ReportExportError):
    code = "report_export_dependency_unavailable"
    status_code = 503


@dataclass(frozen=True)
class ExportedReport:
    content: Any
    format: str
    media_type: str
    extension: str


def normalize_report_format(value: str) -> str:
    """Normalize the public format aliases while rejecting unknown formats."""
    normalized = str(value or "md").strip().lower().lstrip(".")
    try:
        return REPORT_FORMATS[normalized]
    except KeyError as exc:
        supported = ", ".join(sorted(set(REPORT_FORMATS.values())))
        raise UnsupportedReportFormatError(
            f"Unsupported report format: {normalized or '<empty>'}.",
            details={"requested_format": normalized, "supported_formats": supported.split(", ")},
        ) from exc


class ReportExporter:
    """
    Multi-format report exporter.
    Generates professional reports in MD, PDF, and DOCX.
    """

    def __init__(self, author: str = "Nexus AI Pentest Agent"):
        self.author = author
        self.generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S UTC")

    @staticmethod
    def _raw_report(report_data: Dict) -> str:
        """Return a persisted report body without silently truncating it."""
        for key in ("report", "raw_report", "report_markdown"):
            value = report_data.get(key)
            if isinstance(value, str) and value.strip():
                return value
        return ""

    def export(self, report_data: Dict, format: str = "md") -> ExportedReport:
        """Render one of the supported formats with a stable response contract."""
        normalized = normalize_report_format(format)
        if normalized == "pdf":
            return ExportedReport(self.to_pdf(report_data), "pdf", "application/pdf", "pdf")
        if normalized == "docx":
            return ExportedReport(
                self.to_docx(report_data),
                "docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "docx",
            )
        return ExportedReport(self.to_markdown(report_data), "md", "text/markdown", "md")

    def to_markdown(self, report_data: Dict, filter_severity: str = "", filter_cwe: str = "") -> str:
        """Generate Markdown report with optional filtering."""
        target = report_data.get("target", "Unknown")
        findings = report_data.get("findings", [])
        phases = report_data.get("phases", {})
        raw_report = self._raw_report(report_data)

        # A worker-generated report is already the authoritative markdown
        # artifact. Preserve it byte-for-byte when no compatibility filters or
        # extra synthesized sections were requested; the old implementation
        # truncated this body to a 2,000-character phase preview in the API.
        if raw_report and not filter_severity and not filter_cwe and not findings and not phases:
            return raw_report

        # ── Apply filters ─────────────────────────────────────────────────────
        if filter_severity:
            findings = [f for f in findings if f.get("severity", "").lower() == filter_severity.lower()]
        if filter_cwe:
            findings = [f for f in findings if filter_cwe.lower() in str(f.get("cwe", "")).lower()]

        lines = []
        lines.append(f"# Penetration Test Report")
        lines.append("---")
        lines.append("")
        lines.append(f"**Target:** `{target}`")
        lines.append(f"**Date:** {self.generated_at}")
        lines.append(f"**Assessor:** {self.author}")
        lines.append(f"**Total Findings:** {len(findings)}")
        if filter_severity:
            lines.append(f"**Filter:** Severity = {filter_severity}")
        if filter_cwe:
            lines.append(f"**Filter:** CWE = {filter_cwe}")
        lines.append("")

        # Executive Summary
        lines.append("## Executive Summary")
        lines.append("")

        severity_count = {"Critical": 0, "High": 0, "Medium": 0, "Low": 0, "Info": 0}
        for f in findings:
            sev = f.get("severity", "Info")
            severity_count[sev] = severity_count.get(sev, 0) + 1

        lines.append("| Severity | Count | Impact |")
        lines.append("|----------|-------|--------|")
        for sev, count in severity_count.items():
            if count > 0:
                impact = {
                    "Critical": "Immediate action required",
                    "High": "Significant security risk",
                    "Medium": "Should be addressed soon",
                    "Low": "Address when convenient",
                    "Info": "No immediate risk"
                }.get(sev, "")
                lines.append(f"| {sev} | {count} | {impact} |")
        lines.append("")

        # Overall risk
        max_sev = "Info"
        for sev in ["Critical", "High", "Medium", "Low"]:
            if severity_count.get(sev, 0) > 0:
                max_sev = sev
                break
        lines.append(f"**Overall Risk Rating:** {max_sev}")
        lines.append("")

        # Phase results
        phase_names = {"recon": "Reconnaissance", "analis": "Vulnerability Analysis", "eksekutor": "Exploitation", "assessor": "Risk Assessment"}
        for phase_key, phase_name in phase_names.items():
            if phase_key in phases and phases[phase_key]:
                lines.append(f"## {phase_name}")
                lines.append("")
                lines.append(phases[phase_key])
                lines.append("")
                lines.append("---")
                lines.append("")

        if raw_report:
            lines.append("## Persisted Report Content")
            lines.append("")
            lines.append(raw_report)
            lines.append("")

        # Findings
        lines.append("## Detailed Findings")
        lines.append("")

        severity_order = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3, "Info": 4}
        sorted_findings = sorted(findings, key=lambda f: severity_order.get(f.get("severity", "Info"), 5))

        for i, finding in enumerate(sorted_findings, 1):
            sev = finding.get("severity", "Info").lower()
            lines.append(f"### Finding {i}: {finding.get('name', 'Unknown')}")
            lines.append("")
            lines.append(f"**Severity:** {finding.get('severity', 'N/A')}")
            if finding.get("cwe"):
                lines.append(f"**CWE-ID:** {finding['cwe']}")
            if finding.get("cvss"):
                lines.append(f"**CVSS:** {finding['cvss']}")
            if finding.get("url"):
                lines.append(f"**URL:** `{finding['url']}`")
            if finding.get("parameter"):
                lines.append(f"**Parameter:** `{finding['parameter']}`")
            lines.append("")
            if finding.get("description"):
                lines.append(f"**Description:** {finding['description']}")
                lines.append("")
            if finding.get("steps"):
                lines.append("**Steps to Reproduce:**")
                lines.append("```")
                lines.append(finding["steps"])
                lines.append("```")
                lines.append("")
            if finding.get("poc"):
                lines.append("**Proof of Concept:**")
                lines.append("```text")
                lines.append(finding["poc"])
                lines.append("```")
                lines.append("")
            lines.append("---")
            lines.append("")

        # Disclaimer
        lines.append("## Disclaimer")
        lines.append("")
        lines.append("This report was generated by Nexus AI Pentest Agent.")
        lines.append(f"*Report generated on {self.generated_at}*")

        return "\n".join(lines)

    def to_pdf(self, report_data: Dict) -> bytes:
        """
        Generate PDF report.
        Uses fpdf2 library for PDF generation.
        """
        try:
            from fpdf import FPDF
        except ImportError as exc:
            raise ReportDependencyError("fpdf2 is not installed.") from exc

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        # Title
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, "Penetration Test Report", ln=True)
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 8, f"Target: {report_data.get('target', 'Unknown')}", ln=True)
        pdf.cell(0, 8, f"Date: {self.generated_at}", ln=True)
        pdf.cell(0, 8, f"Assessor: {self.author}", ln=True)
        pdf.ln(5)

        # Executive Summary
        pdf.set_font("Helvetica", "B", 14)
        pdf.cell(0, 10, "Executive Summary", ln=True)
        pdf.set_font("Helvetica", "", 10)

        findings = report_data.get("findings", [])
        severity_count = {"Critical": 0, "High": 0, "Medium": 0, "Low": 0, "Info": 0}
        for f in findings:
            sev = f.get("severity", "Info")
            severity_count[sev] = severity_count.get(sev, 0) + 1

        for sev, count in severity_count.items():
            if count > 0:
                pdf.cell(0, 6, f"{sev}: {count}", ln=True)
        pdf.ln(5)

        raw_report = self._raw_report(report_data)
        if raw_report:
            pdf.set_font("Helvetica", "B", 14)
            pdf.cell(0, 10, "Persisted Report Content", ln=True)
            pdf.set_font("Helvetica", "", 9)
            for line in raw_report.splitlines() or [raw_report]:
                # fpdf2's default Helvetica output is latin-1. Replacing only
                # unsupported glyphs keeps export deterministic for existing
                # reports instead of failing the whole format conversion.
                safe_line = str(line).encode("latin-1", "replace").decode("latin-1")
                # fpdf2 advances the cursor to the right margin after a
                # multi_cell call. Reset it before every persisted-report
                # line; otherwise the second line receives zero available
                # width and raises "Not enough horizontal space".
                pdf.set_x(pdf.l_margin)
                pdf.multi_cell(0, 5, safe_line or " ")
            pdf.ln(4)

        # Findings
        pdf.set_font("Helvetica", "B", 14)
        pdf.cell(0, 10, "Detailed Findings", ln=True)
        pdf.set_font("Helvetica", "", 10)

        for i, finding in enumerate(findings, 1):
            if pdf.get_y() > 250:
                pdf.add_page()

            pdf.set_font("Helvetica", "B", 11)
            pdf.cell(0, 8, f"Finding {i}: {finding.get('name', 'Unknown')}", ln=True)
            pdf.set_font("Helvetica", "", 10)
            pdf.cell(0, 6, f"Severity: {finding.get('severity', 'N/A')}", ln=True)
            if finding.get("cwe"):
                pdf.cell(0, 6, f"CWE-ID: {finding['cwe']}", ln=True)
            if finding.get("cvss"):
                pdf.cell(0, 6, f"CVSS: {finding['cvss']}", ln=True)
            if finding.get("description"):
                pdf.multi_cell(0, 6, f"Description: {finding['description']}")
            pdf.ln(3)

        # Footer
        pdf.ln(10)
        pdf.set_font("Helvetica", "I", 8)
        pdf.cell(0, 6, f"Report generated on {self.generated_at}", ln=True)

        output = pdf.output(dest="S")
        if isinstance(output, (bytes, bytearray)):
            return bytes(output)
        return str(output).encode("latin-1", "replace")

    def to_docx(self, report_data: Dict) -> bytes:
        """
        Generate DOCX report.
        Uses python-docx library.
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except ImportError as exc:
            raise ReportDependencyError("python-docx is not installed.") from exc

        doc = Document()

        # Title
        title = doc.add_heading("Penetration Test Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        doc.add_paragraph(f"Target: {report_data.get('target', 'Unknown')}")
        doc.add_paragraph(f"Date: {self.generated_at}")
        doc.add_paragraph(f"Assessor: {self.author}")
        doc.add_paragraph("")

        raw_report = self._raw_report(report_data)
        if raw_report:
            doc.add_heading("Persisted Report Content", level=1)
            doc.add_paragraph(raw_report)
            doc.add_paragraph("")

        # Executive Summary
        doc.add_heading("Executive Summary", level=1)

        findings = report_data.get("findings", [])
        severity_count = {"Critical": 0, "High": 0, "Medium": 0, "Low": 0, "Info": 0}
        for f in findings:
            sev = f.get("severity", "Info")
            severity_count[sev] = severity_count.get(sev, 0) + 1

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        hdr[0].text = "Severity"
        hdr[1].text = "Count"
        for sev, count in severity_count.items():
            if count > 0:
                row = table.add_row().cells
                row[0].text = sev
                row[1].text = str(count)

        doc.add_paragraph("")

        # Findings
        doc.add_heading("Detailed Findings", level=1)

        for i, finding in enumerate(findings, 1):
            doc.add_heading(f"Finding {i}: {finding.get('name', 'Unknown')}", level=2)
            doc.add_paragraph(f"Severity: {finding.get('severity', 'N/A')}")
            if finding.get("cwe"):
                doc.add_paragraph(f"CWE-ID: {finding['cwe']}")
            if finding.get("cvss"):
                doc.add_paragraph(f"CVSS: {finding['cvss']}")
            if finding.get("description"):
                doc.add_paragraph(f"Description: {finding['description']}")
            if finding.get("steps"):
                doc.add_heading("Steps to Reproduce", level=3)
                doc.add_paragraph(finding["steps"])
            if finding.get("poc"):
                doc.add_heading("Proof of Concept", level=3)
                doc.add_paragraph(finding["poc"])
            doc.add_paragraph("")

        # Footer
        doc.add_paragraph("")
        doc.add_paragraph(f"Report generated on {self.generated_at}")

        # Save to bytes
        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()


# Global instance
report_exporter = ReportExporter()
